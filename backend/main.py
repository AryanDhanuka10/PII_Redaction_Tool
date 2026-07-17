"""
main.py — FastAPI backend for the PII Redaction web app.

Endpoints:
  POST /api/redact          upload a .docx -> processes it, returns a job_id + summary report
  GET  /api/download/{id}   download the redacted .docx for that job
  GET  /api/report/{id}     full JSON match log for that job (every value found + its replacement)
  GET  /api/health          liveness check

Design note on "evaluation": for an arbitrary document a user uploads, there's no
ground truth (we don't know in advance where the real PII is), so we can't compute
true precision/recall the way the offline evaluation report does. Instead this API
returns a "redaction summary report": counts per PII type, sample before->after
pairs so a human can spot-check, and a short list of matches flagged as more likely
to be false positives (short/ambiguous spans), so the person reviewing can decide
whether to trust them. This is explained in the frontend and in the README.
"""

import io
import json
import shutil
import uuid
from pathlib import Path

import docx
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from pii_engine import PIIMapper, redact_text

app = FastAPI(title="PII Redaction API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # demo only — restrict this in production
    allow_methods=["*"],
    allow_headers=["*"],
)

JOBS_DIR = Path(__file__).parent / "jobs"
JOBS_DIR.mkdir(exist_ok=True)

# spans this short/ambiguous are worth a human double-checking
REVIEW_FLAG_MAX_LEN = 3


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def process_paragraphs(paragraphs, mapper, log):
    for p in paragraphs:
        original = p.text
        if not original.strip():
            continue
        redacted, matches = redact_text(original, mapper)
        if matches:
            for m in matches:
                log.append({
                    "type": m.label,
                    "original": m.text,
                    "replacement": mapper.fake_for(m.label, m.text),
                    "review_flag": len(m.text.strip()) <= REVIEW_FLAG_MAX_LEN,
                })
            set_paragraph_text(p, redacted)


def process_tables(tables, mapper, log):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, mapper, log)
                if cell.tables:
                    process_tables(cell.tables, mapper, log)


def build_summary(log: list) -> dict:
    by_type = {}
    samples = {}
    flagged = []
    for entry in log:
        t = entry["type"]
        by_type[t] = by_type.get(t, 0) + 1
        samples.setdefault(t, [])
        if len(samples[t]) < 5:
            samples[t].append({"original": entry["original"], "replacement": entry["replacement"]})
        if entry["review_flag"] and len(flagged) < 25:
            flagged.append(entry)

    return {
        "total_redacted": len(log),
        "by_type": by_type,
        "samples": samples,
        "flagged_for_review": flagged,
        "flagged_count": sum(1 for e in log if e["review_flag"]),
    }


@app.get("/api/health")
def health():
    return {"status": "ok"}


@app.post("/api/redact")
async def redact_document(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx files are supported.")

    job_id = str(uuid.uuid4())
    job_dir = JOBS_DIR / job_id
    job_dir.mkdir(parents=True, exist_ok=True)

    input_path = job_dir / "input.docx"
    output_path = job_dir / "redacted.docx"
    log_path = job_dir / "log.json"

    contents = await file.read()
    input_path.write_bytes(contents)

    try:
        document = docx.Document(str(input_path))
    except Exception as e:
        shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(400, f"Could not read this file as a .docx: {e}")

    mapper = PIIMapper()
    log = []

    process_paragraphs(document.paragraphs, mapper, log)
    process_tables(document.tables, mapper, log)
    for section in document.sections:
        process_paragraphs(section.header.paragraphs, mapper, log)
        process_paragraphs(section.footer.paragraphs, mapper, log)

    document.save(str(output_path))
    log_path.write_text(json.dumps(log, indent=2, ensure_ascii=False))

    summary = build_summary(log)
    summary["job_id"] = job_id
    summary["original_filename"] = file.filename
    return summary


@app.get("/api/download/{job_id}")
def download(job_id: str):
    output_path = JOBS_DIR / job_id / "redacted.docx"
    if not output_path.exists():
        raise HTTPException(404, "Job not found or file not ready.")
    return FileResponse(
        str(output_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="redacted.docx",
    )


@app.get("/api/report/{job_id}")
def report(job_id: str):
    log_path = JOBS_DIR / job_id / "log.json"
    if not log_path.exists():
        raise HTTPException(404, "Job not found.")
    return json.loads(log_path.read_text())
