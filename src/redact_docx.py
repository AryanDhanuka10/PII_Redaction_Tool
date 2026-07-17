"""
redact_docx.py
---------------
CLI entry point: reads an input .docx, redacts PII paragraph-by-paragraph
(including inside tables, and nested tables), writes a redacted .docx,
and dumps a JSON log of every match found (used later for evaluation).

Usage:
    python redact_docx.py <input.docx> <output.docx> <matches_log.json>
"""

import sys
import json
import docx
from pii_engine import PIIMapper, redact_text


def set_paragraph_text(paragraph, new_text):
    """Replace a paragraph's visible text while keeping its first run's
    formatting (font, bold, size, ...). Extra runs are cleared."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def process_paragraphs(paragraphs, mapper, log):
    for p in paragraphs:
        original = p.text
        if not original.strip():
            continue
        redacted, matches = redact_text(original, mapper)
        if matches:
            for m in matches:
                log.append({
                    "type": m.label,
                    "original": m.text,
                    "replacement": mapper.fake_for(m.label, m.text),
                })
            set_paragraph_text(p, redacted)


def process_tables(tables, mapper, log):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, mapper, log)
                if cell.tables:
                    process_tables(cell.tables, mapper, log)  # nested tables


def main():
    if len(sys.argv) != 4:
        print("Usage: python redact_docx.py <input.docx> <output.docx> <matches_log.json>")
        sys.exit(1)

    in_path, out_path, log_path = sys.argv[1], sys.argv[2], sys.argv[3]

    document = docx.Document(in_path)
    mapper = PIIMapper()
    log = []

    process_paragraphs(document.paragraphs, mapper, log)
    process_tables(document.tables, mapper, log)

    # headers / footers, if any
    for section in document.sections:
        process_paragraphs(section.header.paragraphs, mapper, log)
        process_paragraphs(section.footer.paragraphs, mapper, log)

    document.save(out_path)

    with open(log_path, "w", encoding="utf-8") as f:
        json.dump(log, f, indent=2, ensure_ascii=False)

    print(f"Done. {len(log)} PII instances redacted.")
    print(f"Redacted file: {out_path}")
    print(f"Match log:     {log_path}")


if __name__ == "__main__":
    main()
